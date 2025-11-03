from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import StreamingResponse, JSONResponse
from io import BytesIO
from docx import Document
import json
import unicodedata
import re

app = FastAPI(title="Leaf Services API")


@app.post("/replace-word/")
async def replace_word(
    file: UploadFile = File(...),
    replacements: str = Form(...)
):
    """
    Recibe un archivo Word (.docx) y reemplaza placeholders con valores.
    Ejemplo replacements:
    {"{{nombre}}": "Jaime", "{{pais}}": "México"}
    """

    # 1️⃣ Validar tipo de archivo
    if not file.filename.endswith(".docx"):
        return JSONResponse(
            status_code=400,
            content={"error": "El archivo debe ser un .docx válido"}
        )

    # 2️⃣ Convertir el texto del form en JSON (diccionario)
    try:
        replacements_dict = json.loads(replacements)
    except json.JSONDecodeError:
        return JSONResponse(
            status_code=400,
            content={"error": "El campo 'replacements' debe ser un JSON válido"}
        )

    # 3️⃣ Leer contenido binario
    content = await file.read()

    # 4️⃣ Cargar documento Word desde memoria
    doc = Document(BytesIO(content))

    # --- 🔁 FUNCIONES DE REEMPLAZO ROBUSTAS ---
    def replace_in_paragraph(paragraph, replacements):
        """Une los runs, reemplaza, y vuelve a escribir el texto en el párrafo."""
        if not paragraph.runs:
            return

        full_text = "".join(run.text for run in paragraph.runs)
        new_text = full_text
        for key, value in replacements.items():
            # Reemplazo directo
            new_text = new_text.replace(key, str(value))
            # Reemplazo flexible (por si hay espacios dentro de {{ }})
            pattern = re.compile(r"\{\{\s*" + re.escape(key.strip("{} ")) + r"\s*\}\}")
            new_text = pattern.sub(str(value), new_text)

        if new_text != full_text:
            # Escribir todo el texto en el primer run
            paragraph.runs[0].text = new_text
            for r in paragraph.runs[1:]:
                r.text = ""

    def replace_in_table(table, replacements):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, replacements)
                # Tablas anidadas
                for nested_table in cell.tables:
                    replace_in_table(nested_table, replacements)

    # --- 🔁 Reemplazar en párrafos principales ---
    for p in doc.paragraphs:
        replace_in_paragraph(p, replacements_dict)

    # --- 🔁 Reemplazar dentro de todas las tablas (recursivo) ---
    for table in doc.tables:
        replace_in_table(table, replacements_dict)

    # --- 🔁 Reemplazar en headers y footers ---
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_in_paragraph(p, replacements_dict)
        for t in section.header.tables:
            replace_in_table(t, replacements_dict)
        for p in section.footer.paragraphs:
            replace_in_paragraph(p, replacements_dict)
        for t in section.footer.tables:
            replace_in_table(t, replacements_dict)

    # 5️⃣ Guardar documento modificado en memoria
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    # 6️⃣ Sanitizar nombre del archivo
    safe_filename = unicodedata.normalize("NFKD", f"modified_{file.filename}")
    safe_filename = safe_filename.encode("ascii", "ignore").decode("ascii")
    safe_filename = re.sub(r"[^A-Za-z0-9_.-]", "_", safe_filename)

    print("\n📄 Archivo procesado correctamente:")
    print(f"Reemplazos aplicados: {replacements_dict}")
    print(f"Archivo devuelto: {safe_filename}")

    # 7️⃣ Devolver archivo como descarga
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={safe_filename}"}
    )
